#!/usr/bin/env python3
"""Safe BatchData 25-city recent-sale validation harness.

The default invocation is a plan-only dry run. Live requests require the
explicit ``--confirm-live`` flag and an API key supplied through
``BATCHDATA_API_KEY`` or ``BATCH_DATA_API_KEY``.

The live workflow independently paginates both supported recent-sale Search
predicates and de-duplicates their record-level union:

* ``searchCriteria.intel.lastSoldDate``
* ``searchCriteria.sale.lastSaleDate``

Persisted artifacts are aggregate and identity-redacted by default. Use
``--sensitive-output`` only for a controlled local investigation; generated
output directories are gitignored, but the operator remains responsible for
protecting those files.
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import json
import math
import os
import sys
import time
import urllib.error
import urllib.request
from dataclasses import dataclass, field
from datetime import date, datetime, timedelta, timezone
from pathlib import Path
from typing import Any, Iterable
from zoneinfo import ZoneInfo


SEARCH_URL = "https://api.batchdata.com/api/v1/property/search"
PHOENIX_TIMEZONE = ZoneInfo("America/Phoenix")
LIVE_PAGE_MAX = 100
DEFAULT_PAGE_SIZE = 100
DEFAULT_MAX_PAGES_PER_STREAM = 5
DEFAULT_ESTIMATED_CREDIT_BUDGET = 500
DEFAULT_WINDOWS = (1, 2, 7, 14)
DEFAULT_MIN_VALUE = 100_000
DEFAULT_SAMPLE_LIMIT = 5


CITIES: tuple[dict[str, Any], ...] = (
    {"city": "Anderson", "state": "SC", "lat": 34.5034, "lng": -82.6501, "region": "Southeast"},
    {"city": "Austin", "state": "TX", "lat": 30.2672, "lng": -97.7431, "region": "Texas"},
    {"city": "Phoenix", "state": "AZ", "lat": 33.4484, "lng": -112.0740, "region": "Southwest"},
    {"city": "Charlotte", "state": "NC", "lat": 35.2271, "lng": -80.8431, "region": "Southeast"},
    {"city": "Indianapolis", "state": "IN", "lat": 39.7684, "lng": -86.1581, "region": "Midwest"},
    {"city": "Atlanta", "state": "GA", "lat": 33.7490, "lng": -84.3880, "region": "Southeast"},
    {"city": "Dallas", "state": "TX", "lat": 32.7767, "lng": -96.7970, "region": "Texas"},
    {"city": "Houston", "state": "TX", "lat": 29.7604, "lng": -95.3698, "region": "Texas"},
    {"city": "San Antonio", "state": "TX", "lat": 29.4241, "lng": -98.4936, "region": "Texas"},
    {"city": "Tampa", "state": "FL", "lat": 27.9506, "lng": -82.4572, "region": "Florida"},
    {"city": "Orlando", "state": "FL", "lat": 28.5383, "lng": -81.3792, "region": "Florida"},
    {"city": "Jacksonville", "state": "FL", "lat": 30.3322, "lng": -81.6557, "region": "Florida"},
    {"city": "Nashville", "state": "TN", "lat": 36.1627, "lng": -86.7816, "region": "Southeast"},
    {"city": "Raleigh", "state": "NC", "lat": 35.7796, "lng": -78.6382, "region": "Southeast"},
    {"city": "Charleston", "state": "SC", "lat": 32.7765, "lng": -79.9311, "region": "Southeast"},
    {"city": "Denver", "state": "CO", "lat": 39.7392, "lng": -104.9903, "region": "Mountain"},
    {"city": "Las Vegas", "state": "NV", "lat": 36.1699, "lng": -115.1398, "region": "Southwest"},
    {"city": "Albuquerque", "state": "NM", "lat": 35.0844, "lng": -106.6504, "region": "Southwest"},
    {"city": "Los Angeles", "state": "CA", "lat": 34.0522, "lng": -118.2437, "region": "West"},
    {"city": "Sacramento", "state": "CA", "lat": 38.5816, "lng": -121.4944, "region": "West"},
    {"city": "Seattle", "state": "WA", "lat": 47.6062, "lng": -122.3321, "region": "West"},
    {"city": "Columbus", "state": "OH", "lat": 39.9612, "lng": -82.9988, "region": "Midwest"},
    {"city": "Kansas City", "state": "MO", "lat": 39.0997, "lng": -94.5786, "region": "Midwest"},
    {"city": "Philadelphia", "state": "PA", "lat": 39.9526, "lng": -75.1652, "region": "Northeast"},
    {"city": "Pittsburgh", "state": "PA", "lat": 40.4406, "lng": -79.9959, "region": "Northeast"},
)


class BudgetExceeded(RuntimeError):
    """Raised before a request that would exceed the local estimate ceiling."""


@dataclass
class ApiResult:
    ok: bool
    status: int
    records: list[dict[str, Any]] = field(default_factory=list)
    total: int | None = None
    elapsed_ms: int = 0
    error_code: str | None = None


@dataclass
class SearchStream:
    predicate: str
    provider_total: int | None
    records: list[dict[str, Any]]
    pages: int
    http_attempts: int
    exhausted: bool
    truncated: bool
    stop_reason: str
    statuses: list[int]


def extract_records(payload: Any) -> list[dict[str, Any]]:
    if not isinstance(payload, dict):
        return []
    candidates = (
        ((payload.get("results") or {}).get("properties") if isinstance(payload.get("results"), dict) else None),
        ((payload.get("results") or {}).get("items") if isinstance(payload.get("results"), dict) else None),
        payload.get("properties"),
        payload.get("items"),
        ((payload.get("data") or {}).get("properties") if isinstance(payload.get("data"), dict) else None),
        payload.get("results"),
    )
    for candidate in candidates:
        if isinstance(candidate, list):
            return [item for item in candidate if isinstance(item, dict)]
    return []


def extract_total(payload: Any) -> int | None:
    if not isinstance(payload, dict):
        return None
    results = payload.get("results") if isinstance(payload.get("results"), dict) else {}
    result_meta = results.get("meta") if isinstance(results.get("meta"), dict) else {}
    meta_results = result_meta.get("results") if isinstance(result_meta.get("results"), dict) else {}
    meta = payload.get("meta") if isinstance(payload.get("meta"), dict) else {}
    values = (
        meta_results.get("resultsFound"),
        results.get("totalRecordCount"),
        payload.get("totalRecordCount"),
        meta.get("totalRecordCount"),
        meta.get("resultsFound"),
    )
    for value in values:
        try:
            if value is not None and value != "":
                parsed = int(value)
                return parsed if parsed >= 0 else None
        except (TypeError, ValueError):
            continue
    return None


class BatchDataClient:
    """Minimal client that never persists provider error bodies or credentials."""

    def __init__(self, api_key: str, estimated_budget: int, timeout_s: int = 25, sleep_s: float = 0.2):
        self._api_key = api_key
        self.estimated_budget = estimated_budget
        self.timeout_s = timeout_s
        self.sleep_s = sleep_s
        self.estimated_credits = 0
        self.http_attempts = 0
        self.status_counts: dict[str, int] = {}

    @property
    def remaining_budget(self) -> int:
        return max(0, self.estimated_budget - self.estimated_credits)

    def post_search(self, payload: dict[str, Any]) -> ApiResult:
        take = int((payload.get("options") or {}).get("take") or 0)
        if take < 0 or take > LIVE_PAGE_MAX:
            raise ValueError(f"options.take must be between 0 and {LIVE_PAGE_MAX}")
        if self.remaining_budget < 1:
            raise BudgetExceeded("Estimated credit budget exhausted before the next Search request.")

        request = urllib.request.Request(
            SEARCH_URL,
            data=json.dumps(payload, separators=(",", ":")).encode("utf-8"),
            headers={
                "Content-Type": "application/json",
                "Authorization": f"Bearer {self._api_key}",
            },
            method="POST",
        )
        started = time.monotonic()
        self.http_attempts += 1
        status = 0
        parsed: Any = {}
        try:
            with urllib.request.urlopen(request, timeout=self.timeout_s) as response:
                status = int(response.status)
                text = response.read().decode("utf-8", errors="replace")
                parsed = json.loads(text) if text else {}
        except urllib.error.HTTPError as exc:
            # Intentionally discard the body: provider errors can echo addresses.
            status = int(exc.code)
        except (urllib.error.URLError, TimeoutError):
            status = 0

        records = extract_records(parsed) if 200 <= status < 300 else []
        # This is a conservative local record-unit estimate, not a billing claim.
        # Failed/count/empty requests reserve one unit; successful record calls use
        # the number of returned rows and never exceed their requested page size.
        estimated_charge = max(1, len(records))
        self.estimated_credits += min(estimated_charge, max(1, take)) if take > 0 else 1
        self.status_counts[str(status)] = self.status_counts.get(str(status), 0) + 1
        if self.sleep_s > 0:
            time.sleep(self.sleep_s)
        return ApiResult(
            ok=200 <= status < 300,
            status=status,
            records=records,
            total=extract_total(parsed) if 200 <= status < 300 else None,
            elapsed_ms=round((time.monotonic() - started) * 1000),
            error_code=None if 200 <= status < 300 else (f"http_{status}" if status else "network_error"),
        )


def square_polygon(city: dict[str, Any], half_size_miles: float = 2.5) -> list[dict[str, float]]:
    lat = float(city["lat"])
    lng = float(city["lng"])
    lat_delta = half_size_miles / 69
    lng_delta = half_size_miles / (69 * math.cos(math.radians(lat)))
    points = (
        (lat + lat_delta, lng - lng_delta),
        (lat + lat_delta, lng + lng_delta),
        (lat - lat_delta, lng + lng_delta),
        (lat - lat_delta, lng - lng_delta),
        (lat + lat_delta, lng - lng_delta),
    )
    return [{"latitude": round(y, 6), "longitude": round(x, 6)} for y, x in points]


def build_search_payload(
    city: dict[str, Any],
    predicate: str,
    min_date: str,
    max_date: str,
    *,
    take: int,
    skip: int,
    qualified: bool = True,
    min_value: int = DEFAULT_MIN_VALUE,
    max_value: int | None = None,
) -> dict[str, Any]:
    if predicate not in {"intel", "sale"}:
        raise ValueError("predicate must be 'intel' or 'sale'")
    if not 0 <= int(take) <= LIVE_PAGE_MAX:
        raise ValueError(f"take must be between 0 and {LIVE_PAGE_MAX}")
    date_range = {"minDate": min_date, "maxDate": max_date}
    criteria: dict[str, Any] = {
        "address": {"geoLocationPolygon": {"geoPoints": square_polygon(city)}},
        predicate: {"lastSoldDate" if predicate == "intel" else "lastSaleDate": date_range},
    }
    if qualified:
        value_range: dict[str, Any] = {"min": max(DEFAULT_MIN_VALUE, int(min_value))}
        if max_value is not None:
            value_range["max"] = int(max_value)
        criteria.update({
            "general": {"standardizedLandUseCode": {"equals": "R2"}},
            "valuation": {"estimatedValue": value_range},
            "listing": {"statusCategory": {"notInList": ["Active", "Pending"]}},
        })
    return {
        "searchCriteria": criteria,
        "options": {
            "take": int(take),
            "skip": max(0, int(skip)),
            "datasets": ["basic"],
        },
    }


def paginate_stream(
    client: BatchDataClient,
    city: dict[str, Any],
    predicate: str,
    min_date: str,
    max_date: str,
    *,
    page_size: int,
    max_pages: int,
    qualified: bool,
    min_value: int,
    max_value: int | None,
) -> SearchStream:
    records: list[dict[str, Any]] = []
    statuses: list[int] = []
    total: int | None = None
    skip = 0
    pages = 0
    exhausted = False
    stop_reason = "max_pages"

    while pages < max_pages:
        if client.remaining_budget <= 0:
            stop_reason = "estimated_budget_exhausted"
            break
        take = min(page_size, LIVE_PAGE_MAX, client.remaining_budget)
        try:
            response = client.post_search(build_search_payload(
                city,
                predicate,
                min_date,
                max_date,
                take=take,
                skip=skip,
                qualified=qualified,
                min_value=min_value,
                max_value=max_value,
            ))
        except BudgetExceeded:
            stop_reason = "estimated_budget_exhausted"
            break
        pages += 1
        statuses.append(response.status)
        if not response.ok:
            stop_reason = response.error_code or "request_failed"
            break
        if total is None:
            total = response.total
        page_records = response.records
        records.extend(page_records)
        skip += len(page_records)
        if not page_records:
            exhausted = total in {None, 0} or skip >= (total or 0)
            stop_reason = "provider_exhausted" if exhausted else "empty_page_before_reported_total"
            break
        if total is not None and skip >= total:
            exhausted = True
            stop_reason = "provider_exhausted"
            break
        if len(page_records) < take:
            exhausted = True
            stop_reason = "short_final_page"
            break

    return SearchStream(
        predicate=predicate,
        provider_total=total,
        records=records,
        pages=pages,
        http_attempts=pages,
        exhausted=exhausted,
        truncated=not exhausted,
        stop_reason=stop_reason,
        statuses=statuses,
    )


def property_object(record: dict[str, Any]) -> dict[str, Any]:
    for candidate in (
        record.get("property"),
        (record.get("result") or {}).get("property") if isinstance(record.get("result"), dict) else None,
        (record.get("response") or {}).get("property") if isinstance(record.get("response"), dict) else None,
        record,
    ):
        if isinstance(candidate, dict):
            return candidate
    return {}


def normalized_address(property_data: dict[str, Any]) -> dict[str, str]:
    address = property_data.get("address") if isinstance(property_data.get("address"), dict) else {}
    return {
        "street": str(address.get("street") or address.get("streetAddress") or address.get("addressLine1") or property_data.get("addressLine1") or "").strip(),
        "city": str(address.get("city") or property_data.get("city") or "").strip(),
        "state": str(address.get("state") or property_data.get("state") or "").strip(),
        "zip": str(address.get("zip") or address.get("zipCode") or property_data.get("zipCode") or "").strip()[:5],
    }


def record_identity(record: dict[str, Any]) -> str:
    property_data = property_object(record)
    identifiers = property_data.get("ids") if isinstance(property_data.get("ids"), dict) else {}
    provider_id = identifiers.get("propertyId") or identifiers.get("id") or property_data.get("propertyId") or property_data.get("id")
    if provider_id:
        return f"provider:{provider_id}"
    address = normalized_address(property_data)
    normalized = "|".join(address.values()).lower()
    if all(address.values()):
        return f"address:{normalized}"
    stable = json.dumps(record, sort_keys=True, default=str, separators=(",", ":"))
    return f"record:{hashlib.sha256(stable.encode('utf-8')).hexdigest()}"


def union_stream_records(intel: SearchStream, sale: SearchStream) -> list[dict[str, Any]]:
    union: dict[str, dict[str, Any]] = {}
    for stream in (intel, sale):
        for record in stream.records:
            identity = record_identity(record)
            entry = union.setdefault(identity, {"identity": identity, "record": record, "sources": set()})
            entry["sources"].add(stream.predicate)
            if stream.predicate == "sale":
                entry["record"] = record
    return sorted(
        union.values(),
        key=lambda item: (
            0 if item["sources"] == {"intel", "sale"} else 1 if "sale" in item["sources"] else 2,
            hashlib.sha256(item["identity"].encode("utf-8")).hexdigest(),
        ),
    )


def first_value(*values: Any) -> Any:
    return next((value for value in values if value not in (None, "")), None)


def safe_sample(union_entry: dict[str, Any], sensitive_output: bool = False) -> dict[str, Any]:
    record = union_entry["record"]
    property_data = property_object(record)
    address = property_data.get("address") if isinstance(property_data.get("address"), dict) else {}
    general = property_data.get("general") if isinstance(property_data.get("general"), dict) else {}
    intel = property_data.get("intel") if isinstance(property_data.get("intel"), dict) else {}
    sale = property_data.get("sale") if isinstance(property_data.get("sale"), dict) else {}
    listing = property_data.get("listing") if isinstance(property_data.get("listing"), dict) else {}
    valuation = property_data.get("valuation") if isinstance(property_data.get("valuation"), dict) else {}
    identity_hash = hashlib.sha256(union_entry["identity"].encode("utf-8")).hexdigest()
    sample: dict[str, Any] = {
        "identity_sha256": identity_hash,
        "sources": sorted(union_entry["sources"]),
        "field_presence": {
            "coordinates": first_value(address.get("latitude"), property_data.get("latitude")) is not None
                and first_value(address.get("longitude"), property_data.get("longitude")) is not None,
            "intel_last_sold_date": intel.get("lastSoldDate") not in (None, ""),
            "sale_last_sale_date": first_value(sale.get("lastSaleDate"), property_data.get("lastSaleDate")) is not None,
            "standardized_land_use": first_value(general.get("standardizedLandUseCode"), property_data.get("standardizedLandUseCode")) is not None,
            "estimated_value": first_value(valuation.get("estimatedValue"), property_data.get("estimatedValue")) is not None,
            "listing_status": first_value(listing.get("statusCategory"), listing.get("status")) is not None,
            "owner": isinstance(property_data.get("owner"), dict) and bool(property_data.get("owner")),
        },
    }
    if sensitive_output:
        identifiers = property_data.get("ids") if isinstance(property_data.get("ids"), dict) else {}
        sample["sensitive"] = {
            "address": normalized_address(property_data),
            "provider_property_id": first_value(identifiers.get("propertyId"), property_data.get("propertyId")),
            "raw_provider_record": record,
        }
    return sample


def run_city_window(
    client: BatchDataClient,
    city: dict[str, Any],
    window_days: int,
    yesterday: date,
    args: argparse.Namespace,
) -> dict[str, Any]:
    min_day = yesterday - timedelta(days=window_days - 1)
    min_date = min_day.isoformat()
    max_date = yesterday.isoformat()
    streams: dict[str, SearchStream] = {}
    for predicate in ("intel", "sale"):
        streams[predicate] = paginate_stream(
            client,
            city,
            predicate,
            min_date,
            max_date,
            page_size=args.page_size,
            max_pages=args.max_pages_per_stream,
            qualified=not args.broad,
            min_value=args.min_value,
            max_value=args.max_value,
        )
    union = union_stream_records(streams["intel"], streams["sale"])
    overlap = sum(1 for item in union if item["sources"] == {"intel", "sale"})
    intel_only = sum(1 for item in union if item["sources"] == {"intel"})
    sale_only = sum(1 for item in union if item["sources"] == {"sale"})
    complete = streams["intel"].exhausted and streams["sale"].exhausted
    return {
        "city": city["city"],
        "state": city["state"],
        "region": city["region"],
        "window_days": window_days,
        "min_date": min_date,
        "max_date": max_date,
        "qualified_route_filters": not args.broad,
        "intel": stream_summary(streams["intel"]),
        "sale": stream_summary(streams["sale"]),
        "union": {
            "complete": complete,
            "sampled_or_exact_count": len(union),
            "count_semantics": "exact" if complete else "lower_bound_truncated",
            "overlap": overlap,
            "intel_only": intel_only,
            "sale_only": sale_only,
            "samples": [safe_sample(item, args.sensitive_output) for item in union[: args.sample_limit]],
        },
    }


def stream_summary(stream: SearchStream) -> dict[str, Any]:
    return {
        "provider_total": stream.provider_total,
        "records_reviewed": len(stream.records),
        "pages": stream.pages,
        "http_attempts": stream.http_attempts,
        "exhausted": stream.exhausted,
        "truncated": stream.truncated,
        "stop_reason": stream.stop_reason,
        "http_statuses": stream.statuses,
    }


def parse_windows(value: str) -> tuple[int, ...]:
    windows: list[int] = []
    for raw in value.split(","):
        try:
            days = int(raw.strip())
        except ValueError as exc:
            raise argparse.ArgumentTypeError(f"Invalid window: {raw}") from exc
        if days <= 0 or days > 365:
            raise argparse.ArgumentTypeError("Windows must be between 1 and 365 days")
        if days not in windows:
            windows.append(days)
    if not windows:
        raise argparse.ArgumentTypeError("At least one window is required")
    return tuple(windows)


def selected_cities(args: argparse.Namespace) -> list[dict[str, Any]]:
    cities = list(CITIES)
    if args.city:
        wanted = {item.strip().lower() for item in args.city.split(",") if item.strip()}
        cities = [item for item in cities if item["city"].lower() in wanted]
        if not cities:
            raise ValueError("--city did not match a configured city")
    if args.limit_cities:
        cities = cities[: args.limit_cities]
    return cities


def build_plan(args: argparse.Namespace) -> dict[str, Any]:
    cities = selected_cities(args)
    windows = parse_windows(args.windows)
    return {
        "mode": "live" if args.confirm_live else "plan_only_no_network",
        "network_requests_made": 0,
        "endpoint": SEARCH_URL,
        "city_count": len(cities),
        "cities": [f"{item['city']}, {item['state']}" for item in cities],
        "windows_days": list(windows),
        "predicates": ["intel.lastSoldDate", "sale.lastSaleDate"],
        "relationship": "independent streams; record-level union after separate pagination",
        "page_size": args.page_size,
        "live_page_max": LIVE_PAGE_MAX,
        "max_pages_per_stream": args.max_pages_per_stream,
        "estimated_request_ceiling_before_budget_stop": len(cities) * len(windows) * 2 * args.max_pages_per_stream,
        "estimated_credit_budget": args.budget,
        "qualified_filters": None if args.broad else {
            "standardized_land_use_code": "R2",
            "estimated_value_min": args.min_value,
            "estimated_value_max": args.max_value,
            "listing_status_excluded": ["Active", "Pending"],
        },
        "privacy": {
            "sensitive_output": args.sensitive_output,
            "default_artifacts_contain_raw_addresses": False,
            "default_artifacts_contain_raw_provider_payloads": False,
        },
        "billing_warning": "The budget is a local record-unit estimate, not a provider billing guarantee. Reconcile exact charges in BatchData.",
    }


def flatten_rows(results: Iterable[dict[str, Any]]) -> list[dict[str, Any]]:
    rows = []
    for result in results:
        rows.append({
            "city": result["city"],
            "state": result["state"],
            "region": result["region"],
            "window_days": result["window_days"],
            "min_date": result["min_date"],
            "max_date": result["max_date"],
            "intel_provider_total": result["intel"]["provider_total"],
            "intel_reviewed": result["intel"]["records_reviewed"],
            "intel_exhausted": result["intel"]["exhausted"],
            "sale_provider_total": result["sale"]["provider_total"],
            "sale_reviewed": result["sale"]["records_reviewed"],
            "sale_exhausted": result["sale"]["exhausted"],
            "union_count": result["union"]["sampled_or_exact_count"],
            "union_count_semantics": result["union"]["count_semantics"],
            "overlap": result["union"]["overlap"],
            "intel_only": result["union"]["intel_only"],
            "sale_only": result["union"]["sale_only"],
        })
    return rows


def write_csv(path: Path, rows: list[dict[str, Any]]) -> None:
    if not rows:
        path.write_text("", encoding="utf-8")
        return
    with path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(rows[0]))
        writer.writeheader()
        writer.writerows(rows)


def build_markdown_report(payload: dict[str, Any]) -> str:
    metadata = payload["metadata"]
    rows = payload["summary_rows"]
    lines = [
        "# BatchData 25-city recent-sale validation",
        "",
        f"Generated: {metadata['completed_at']}",
        "",
        "## Contract",
        "",
        "Intel and Sale are independent predicates on the same Property Search endpoint. Each stream was paginated separately with a maximum page size of 100, then de-duplicated into a record-level union.",
        "",
        f"Estimated local record units: {metadata['estimated_credits_used']}/{metadata['estimated_credit_budget']}. Exact provider billing was not inferred.",
        "",
        "Artifacts are identity-redacted unless `sensitive_output` is true in metadata.",
        "",
        "## Results",
        "",
        "| City | Window | Intel total | Sale total | Union | Semantics | Intel-only | Sale-only |",
        "|---|---:|---:|---:|---:|---|---:|---:|",
    ]
    for row in rows:
        lines.append(
            f"| {row['city']}, {row['state']} | {row['window_days']}d | "
            f"{row['intel_provider_total'] if row['intel_provider_total'] is not None else 'n/a'} | "
            f"{row['sale_provider_total'] if row['sale_provider_total'] is not None else 'n/a'} | "
            f"{row['union_count']} | {row['union_count_semantics']} | {row['intel_only']} | {row['sale_only']} |"
        )
    lines.extend([
        "",
        "## Limitations",
        "",
        "- These are BatchData candidates, not recall against an independent recorder or MLS truth set.",
        "- Provider filter acceptance is not an independent correctness audit of R2, valuation, listing, or sale dates.",
        "- Any row marked `lower_bound_truncated` stopped before both streams were exhausted.",
        "- API attempts, returned rows, local estimated record units, and billable credits are different quantities.",
        "",
    ])
    return "\n".join(lines)


def build_docx_report(path: Path, markdown_payload: dict[str, Any]) -> None:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("DOCX output requires the optional python-docx package") from exc
    document = Document()
    document.add_heading("BatchData 25-city recent-sale validation", 0)
    metadata = markdown_payload["metadata"]
    document.add_paragraph(
        "Intel and Sale were paginated independently with pages no larger than 100 records, then de-duplicated into a record-level union."
    )
    document.add_paragraph(
        f"Estimated local record units: {metadata['estimated_credits_used']}/{metadata['estimated_credit_budget']}. Exact provider billing was not inferred."
    )
    table = document.add_table(rows=1, cols=8)
    headers = ["City", "Window", "Intel", "Sale", "Union", "Semantics", "Intel-only", "Sale-only"]
    for index, value in enumerate(headers):
        table.rows[0].cells[index].text = value
    for row in markdown_payload["summary_rows"]:
        cells = table.add_row().cells
        values = (
            f"{row['city']}, {row['state']}",
            f"{row['window_days']}d",
            row["intel_provider_total"],
            row["sale_provider_total"],
            row["union_count"],
            row["union_count_semantics"],
            row["intel_only"],
            row["sale_only"],
        )
        for index, value in enumerate(values):
            cells[index].text = "n/a" if value is None else str(value)
    document.add_heading("Limitations", level=1)
    for limitation in (
        "Provider candidates are not an external recorder/MLS recall denominator.",
        "Filter acceptance does not independently prove provider classification correctness.",
        "Truncated unions are lower bounds.",
        "Local estimated record units are not a provider billing statement.",
    ):
        document.add_paragraph(limitation, style="List Bullet")
    document.save(path)


def parse_args(argv: list[str] | None = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Safely validate BatchData Intel and Sale Search coverage across 25 cities.")
    parser.add_argument("--confirm-live", action="store_true", help="Required before any network request is made.")
    parser.add_argument("--budget", type=int, default=DEFAULT_ESTIMATED_CREDIT_BUDGET, help="Local estimated record-unit ceiling (default: 500).")
    parser.add_argument("--page-size", type=int, default=DEFAULT_PAGE_SIZE, help="Records per Search page; maximum 100.")
    parser.add_argument("--max-pages-per-stream", type=int, default=DEFAULT_MAX_PAGES_PER_STREAM)
    parser.add_argument("--windows", default=",".join(map(str, DEFAULT_WINDOWS)), help="Comma-separated completed-day windows.")
    parser.add_argument("--as-of", help="America/Phoenix YYYY-MM-DD; windows end on the prior day.")
    parser.add_argument("--city", help="Comma-separated configured city names.")
    parser.add_argument("--limit-cities", type=int, default=0)
    parser.add_argument("--min-value", type=int, default=DEFAULT_MIN_VALUE)
    parser.add_argument("--max-value", type=int)
    parser.add_argument("--broad", action="store_true", help="Omit R2/value/listing filters for a broad diagnostic run.")
    parser.add_argument("--sample-limit", type=int, default=DEFAULT_SAMPLE_LIMIT)
    parser.add_argument("--sensitive-output", action="store_true", help="Persist sampled addresses, provider ids, and raw sample records. Keep output local and protected.")
    parser.add_argument("--output-dir", default="batchdata_results")
    parser.add_argument("--docx", action="store_true", help="Also create a DOCX report (requires python-docx).")
    parser.add_argument("--timeout", type=int, default=25)
    parser.add_argument("--sleep", type=float, default=0.2)
    args = parser.parse_args(argv)
    if args.budget <= 0:
        parser.error("--budget must be positive")
    if not 1 <= args.page_size <= LIVE_PAGE_MAX:
        parser.error(f"--page-size must be between 1 and {LIVE_PAGE_MAX}")
    if args.max_pages_per_stream <= 0:
        parser.error("--max-pages-per-stream must be positive")
    if args.limit_cities < 0 or args.limit_cities > len(CITIES):
        parser.error(f"--limit-cities must be between 0 and {len(CITIES)}")
    if args.min_value < DEFAULT_MIN_VALUE:
        parser.error(f"--min-value cannot be below the required ${DEFAULT_MIN_VALUE:,} floor")
    if args.max_value is not None and args.max_value < args.min_value:
        parser.error("--max-value cannot be below --min-value")
    if args.sample_limit < 0 or args.sample_limit > 25:
        parser.error("--sample-limit must be between 0 and 25")
    parse_windows(args.windows)
    return args


def parse_as_of(value: str | None) -> date:
    if not value:
        return datetime.now(PHOENIX_TIMEZONE).date()
    try:
        return date.fromisoformat(value)
    except ValueError as exc:
        raise ValueError("--as-of must be YYYY-MM-DD") from exc


def main(argv: list[str] | None = None) -> int:
    args = parse_args(argv)
    plan = build_plan(args)
    if not args.confirm_live:
        print(json.dumps(plan, indent=2))
        print("\nPLAN ONLY: no BatchData request was made. Re-run with --confirm-live to execute.", file=sys.stderr)
        return 0

    api_key = os.environ.get("BATCHDATA_API_KEY") or os.environ.get("BATCH_DATA_API_KEY")
    if not api_key:
        print("BATCHDATA_API_KEY or BATCH_DATA_API_KEY must be set for --confirm-live.", file=sys.stderr)
        return 2
    if args.sensitive_output:
        print("WARNING: --sensitive-output will persist sampled addresses, provider ids, and raw provider records.", file=sys.stderr)

    as_of = parse_as_of(args.as_of)
    yesterday = as_of - timedelta(days=1)
    cities = selected_cities(args)
    windows = parse_windows(args.windows)
    client = BatchDataClient(api_key, args.budget, timeout_s=args.timeout, sleep_s=args.sleep)
    results: list[dict[str, Any]] = []
    started_at = datetime.now(timezone.utc).isoformat()
    budget_stopped = False

    for city in cities:
        for window_days in windows:
            if client.remaining_budget <= 0:
                budget_stopped = True
                break
            print(
                f"[{len(results) + 1}/{len(cities) * len(windows)}] {city['city']}, {city['state']} {window_days}d "
                f"(estimated budget remaining: {client.remaining_budget})",
                file=sys.stderr,
            )
            result = run_city_window(client, city, window_days, yesterday, args)
            results.append(result)
            if client.remaining_budget <= 0:
                budget_stopped = True
                break
        if budget_stopped:
            break

    completed_at = datetime.now(timezone.utc).isoformat()
    rows = flatten_rows(results)
    payload = {
        "schema_version": 1,
        "metadata": {
            "started_at": started_at,
            "completed_at": completed_at,
            "as_of_date": as_of.isoformat(),
            "yesterday_date": yesterday.isoformat(),
            "timezone": "America/Phoenix",
            "city_count_requested": len(cities),
            "city_window_rows_completed": len(results),
            "windows_days": list(windows),
            "estimated_credit_budget": args.budget,
            "estimated_credits_used": client.estimated_credits,
            "budget_stopped": budget_stopped,
            "http_attempts": client.http_attempts,
            "http_status_counts": client.status_counts,
            "page_size": args.page_size,
            "live_page_max": LIVE_PAGE_MAX,
            "max_pages_per_stream": args.max_pages_per_stream,
            "sensitive_output": args.sensitive_output,
            "raw_addresses_persisted": args.sensitive_output,
            "raw_provider_payload_samples_persisted": args.sensitive_output,
            "billing_status": "unverified_confirm_in_batchdata_dashboard",
        },
        "contract": {
            "endpoint": SEARCH_URL,
            "datasets_location": "options.datasets",
            "datasets": ["basic"],
            "predicates": ["searchCriteria.intel.lastSoldDate", "searchCriteria.sale.lastSaleDate"],
            "union": "separately paginated, record-level de-duplication",
            "route_filters": None if args.broad else {
                "general.standardizedLandUseCode.equals": "R2",
                "valuation.estimatedValue.min": args.min_value,
                "valuation.estimatedValue.max": args.max_value,
                "listing.statusCategory.notInList": ["Active", "Pending"],
            },
        },
        "results": results,
        "summary_rows": rows,
        "limitations": [
            "Provider candidate coverage is not recall against an external recorder or MLS truth set.",
            "Filter acceptance is not independent proof of provider classification correctness.",
            "A non-exhausted stream makes its union a lower bound.",
            "The local record-unit budget is not a provider billing guarantee.",
        ],
    }

    slug = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%SZ")
    output_dir = Path(args.output_dir) / slug
    output_dir.mkdir(parents=True, exist_ok=True)
    json_path = output_dir / "batchdata_25_city_validation.json"
    csv_path = output_dir / "batchdata_25_city_summary.csv"
    markdown_path = output_dir / "batchdata_25_city_report.md"
    json_path.write_text(json.dumps(payload, indent=2), encoding="utf-8")
    write_csv(csv_path, rows)
    markdown_path.write_text(build_markdown_report(payload), encoding="utf-8")
    if args.docx:
        build_docx_report(output_dir / "batchdata_25_city_report.docx", payload)

    print(f"Output directory: {output_dir}")
    print(f"Estimated record units: {client.estimated_credits}/{client.estimated_budget}")
    print(f"HTTP attempts: {client.http_attempts}")
    print(f"JSON: {json_path}")
    print(f"CSV: {csv_path}")
    print(f"Markdown: {markdown_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
